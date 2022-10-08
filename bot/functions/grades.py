import asyncio
import subprocess
import os
import shutil

from aiogram import types
from docx import Document

from bot.keyboards.default import add_delete_button


async def to_pdf(filepath, path):
    return subprocess.call(
        ['soffice',
        '--headless',
        '--convert-to',
        'pdf',
        '--outdir',
        path,
        filepath]
    ), filepath


async def send_files(message, dir):
    media = types.MediaGroup()
    media2 = types.MediaGroup()
    files = []
    i = 0
    for filename in os.listdir(os.getcwd()+f'/{dir}'):
        files.append(
            open(os.path.join(os.getcwd()+f'/{dir}', filename), 'rb'))
        if i < 10:
            media.attach_document(files[i])
        elif i < 20:
            media2.attach_document(files[i])
        i += 1
    if i > 0:
        await message.bot.send_media_group(message.chat.id, media=media, disable_notification=True)
        if i > 10:
            await message.bot.send_media_group(message.chat.id, media=media2, disable_notification=True)
    else:
        await message.answer('No files', reply_markup=add_delete_button())
    for item in files:
        item.close()


async def set_grades_course(document, course, is_active_only):
    if is_active_only:
        if course['active'] == 0:
            return []
    course_name = course['name']
    course_id = course['id']
    save_name = None
    save_grade = None

    document.add_heading(f'{course_name}', 0)

    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Grade item'
    hdr_cells[1].text = 'Percentage'
    for grade in course['grades']:
        percentage = course['grades'][grade]['percentage']
        name = course['grades'][grade]['name']
        row_cells = table.add_row().cells
        try:
            name = name.replace('', '')
            row_cells[0].text = name
        except:
            row_cells[0].text = name.split('.')[0]
        row_cells[1].text = percentage
        if course['grades'][grade]['name']=='Course total':
            save_name = course_name
            save_grade = course['grades'][grade]['percentage']
    document.add_page_break()

    return [save_name, save_grade]


async def set_total(group, barcode):
    document = Document()
    document.add_heading('Total Grades', 0)

    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Cource'
    hdr_cells[1].text = 'Grade'
    for course in group:
        name = course[0]
        grade = course[1]
        if name is None: continue
        row_cells = table.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = grade

    document.add_page_break()

    document.save(f'{barcode}_grades/TOTAL.docx')


async def local_grades(user, message, is_active_only):
    barcode = user['barcode']
    if is_active_only:
        path = f'{barcode}_grades_active'
    else:
        path = f'{barcode}_grades'

    if not os.path.exists(path):
        os.makedirs(path)
    
    document = Document()
    group = await asyncio.gather(*[set_grades_course(document, user['courses'][key], is_active_only) for key in user['courses']])
    if not is_active_only:
        await set_total(group, barcode)
    document.save(f'{path}/GRADES.docx')

    await message.bot.send_chat_action(message.chat.id, types.ChatActions.UPLOAD_DOCUMENT)
    processes = await asyncio.gather(*[to_pdf(f'{path}/{filename}', path) for filename in os.listdir(os.getcwd()+f'/{path}')])
    for proc, filepath in processes:
        os.remove(filepath)

    await send_files(message, path)
    try:
        shutil.rmtree(path)
    except:
        pass
